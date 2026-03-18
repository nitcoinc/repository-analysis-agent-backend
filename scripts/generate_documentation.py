"""
Script to generate professional Word documents for the Codebase Analysis Agent project.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def setup_document_style(doc):
    """Setup document styles."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Create heading styles
    heading1 = doc.styles['Heading 1']
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    
    heading2 = doc.styles['Heading 2']
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 102, 204)
    
    heading3 = doc.styles['Heading 3']
    heading3.font.size = Pt(12)
    heading3.font.bold = True

def add_title_page(doc, title, subtitle, version="1.0.0"):
    """Add a professional title page."""
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Version
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run(f"Version {version}")
    version_run.font.size = Pt(12)
    version_run.font.italic = True
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents placeholder."""
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('(To be updated automatically in Word)')
    doc.add_page_break()

def create_process_flow_document():
    """Create High-Level Process Flow document."""
    doc = Document()
    setup_document_style(doc)
    
    # Title Page
    add_title_page(doc, "High-Level Process Flow", 
                   "Codebase Analysis Agent System", "1.0.0")
    
    # Table of Contents
    add_table_of_contents(doc)
    
    # Overview
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'The Codebase Analysis Agent system follows a multi-agent orchestration pattern '
        'where specialized AI agents work together to analyze codebases, map dependencies, '
        'and generate insights. This document describes the high-level process flows for '
        'all major operations in the system.'
    )
    
    # Main Process Flow
    doc.add_heading('2. Main Process Flow', 1)
    
    doc.add_heading('2.1 Phase 1: Repository Acquisition & Setup', 2)
    doc.add_paragraph(
        'The system begins by acquiring the repository from one of three sources:'
    )
    p = doc.add_paragraph('• ', style='List Bullet')
    p.add_run('GitHub API').bold = True
    p.add_run(' - Direct access via GitHub API with authentication token')
    
    p = doc.add_paragraph('• ', style='List Bullet')
    p.add_run('Git Clone').bold = True
    p.add_run(' - Clone from any Git URL (public or private with credentials)')
    
    p = doc.add_paragraph('• ', style='List Bullet')
    p.add_run('Local Path').bold = True
    p.add_run(' - Use existing local repository on the file system')
    
    doc.add_paragraph(
        'Once the repository is acquired, a unique repository ID is generated and an '
        'analysis run is created. The analysis is queued as a background task to allow '
        'non-blocking API responses.'
    )
    
    doc.add_heading('2.2 Phase 2: Agent Orchestration Workflow', 2)
    doc.add_paragraph(
        'The orchestrator coordinates a sequence of specialized agents, each responsible '
        'for a specific aspect of codebase analysis:'
    )
    
    agents = [
        ('Planning Agent', [
            'Analyzes codebase structure',
            'Identifies programming languages',
            'Creates execution plan with estimated time',
            'Determines analysis scope and priorities'
        ]),
        ('Code Browser Agent', [
            'Recursively explores all code files',
            'Parses code files using language-specific parsers',
            'Extracts code elements (functions, classes, methods)',
            'Collects import statements and dependencies'
        ]),
        ('Dependency Mapper Agent', [
            'Identifies service boundaries automatically',
            'Maps service-to-service dependencies',
            'Detects API endpoint calls',
            'Finds database connections',
            'Identifies message queue usage',
            'Builds Neo4j dependency graph'
        ]),
        ('Documentation Agent', [
            'Generates service documentation using OpenAI LLM',
            'Creates API specifications',
            'Documents functionality and purpose',
            'Maintains documentation versioning'
        ]),
        ('Impact Agent', [
            'Prepares impact analysis capabilities',
            'Builds dependency graph cache',
            'Sets up impact analysis infrastructure'
        ]),
        ('Human Review Agent', [
            'Identifies ambiguous cases requiring human input',
            'Creates checkpoints for unclear dependencies',
            'Flags legacy code patterns',
            'Pauses workflow when checkpoints are created'
        ])
    ]
    
    for agent_name, tasks in agents:
        doc.add_heading(f'{agent_name}', 3)
        for task in tasks:
            doc.add_paragraph(f'• {task}', style='List Bullet')
    
    doc.add_heading('2.3 Phase 3: Impact Analysis Flow', 2)
    doc.add_paragraph(
        'When a user requests impact analysis for a proposed change:'
    )
    flow_steps = [
        'User provides change description and optionally affected files/services',
        'Impact Engine receives the request',
        'System traverses the dependency graph in Neo4j to find affected services',
        'Impact scores are calculated based on dependency depth and change type',
        'Risk level is determined (low, medium, high, critical)',
        'Recommendations are generated based on impact and risk',
        'Comprehensive impact report is returned to the user'
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('2.4 Phase 4: Human-in-the-Loop Flow', 2)
    doc.add_paragraph(
        'When an agent encounters ambiguous cases that require human judgment:'
    )
    hitl_steps = [
        'Agent detects ambiguity (unclear dependencies, service boundaries, etc.)',
        'Checkpoint is created with status "pending"',
        'Workflow execution is paused',
        'User is notified via the frontend interface',
        'User reviews the checkpoint with context and options',
        'User provides a response or clarification',
        'Checkpoint status is updated to "resolved"',
        'Workflow execution resumes from the checkpoint'
    ]
    for i, step in enumerate(hitl_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    # Data Flow
    doc.add_heading('3. Data Flow', 1)
    
    doc.add_heading('3.1 Repository Analysis Data Flow', 2)
    doc.add_paragraph(
        'The data flows through the following stages during repository analysis:'
    )
    data_flow = [
        ('Repository Files', 'Raw source code files from the repository'),
        ('Code Parser Service', 'Parses files and extracts code elements, imports, dependencies'),
        ('Dependency Analyzer', 'Identifies service dependencies, API calls, database connections'),
        ('Graph Service (Neo4j)', 'Stores dependency relationships in graph database'),
        ('PostgreSQL', 'Stores metadata: services, documentation, analysis runs')
    ]
    for stage, description in data_flow:
        p = doc.add_paragraph()
        p.add_run(f'{stage}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('3.2 Frontend Data Flow', 2)
    doc.add_paragraph(
        'User interactions flow through the following layers:'
    )
    frontend_flow = [
        'User Interface (React Components)',
        'React Query / API Client (Data fetching)',
        'FastAPI Backend (REST API)',
        'Services / Graph Service / Impact Engine (Business Logic)',
        'Neo4j / PostgreSQL / Redis Cache (Data Storage)',
        'Response Data (JSON)',
        'React Components (Rendering)',
        'Visualization (React Flow, Charts)'
    ]
    for step in frontend_flow:
        doc.add_paragraph(f'• {step}', style='List Bullet')
    
    # State Management
    doc.add_heading('4. State Management', 1)
    doc.add_paragraph(
        'AgentState maintains execution context throughout the workflow:'
    )
    doc.add_paragraph('• Data: Current analysis data and results', style='List Bullet')
    doc.add_paragraph('• History: Execution history for each agent', style='List Bullet')
    doc.add_paragraph('• Checkpoints: Human review checkpoints', style='List Bullet')
    doc.add_paragraph(
        'State is passed sequentially from one agent to the next, with each agent '
        'updating the state with its results. The final state is stored in the database.'
    )
    
    # Error Handling
    doc.add_heading('5. Error Handling Flow', 1)
    error_steps = [
        'Agent execution encounters an error',
        'Exception is caught and logged with context',
        'Error metrics are updated in Prometheus',
        'State is updated with status "failed"',
        'Error details are stored for debugging',
        'User is notified via API response with error details'
    ]
    for i, step in enumerate(error_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    # Caching
    doc.add_heading('6. Caching Flow', 1)
    doc.add_paragraph(
        'The system uses Redis for caching to improve performance:'
    )
    doc.add_paragraph(
        '1. Request for parsed code or dependency graph is received',
        style='List Number'
    )
    doc.add_paragraph(
        '2. System checks Redis cache for existing data',
        style='List Number'
    )
    doc.add_paragraph(
        '3. If cache hit: Return cached data immediately',
        style='List Number'
    )
    doc.add_paragraph(
        '4. If cache miss: Parse/analyze code, store in cache with TTL, return data',
        style='List Number'
    )
    
    # Real-time Updates
    doc.add_heading('7. Real-time Status Updates', 1)
    doc.add_paragraph(
        'Status updates are provided through polling (with future WebSocket support):'
    )
    doc.add_paragraph('• Background task updates state after each agent execution', style='List Bullet')
    doc.add_paragraph('• Frontend polls status endpoint every 5 seconds', style='List Bullet')
    doc.add_paragraph('• UI updates with current progress and status', style='List Bullet')
    doc.add_paragraph('• Future: WebSocket support for push-based updates', style='List Bullet')
    
    # Save document
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 
                               'docs', 'High-Level_Process_Flow.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

def create_requirements_document():
    """Create Requirements Document."""
    doc = Document()
    setup_document_style(doc)
    
    # Title Page
    add_title_page(doc, "Requirements Document", 
                   "Codebase Analysis Agent System", "1.0.0")
    
    # Table of Contents
    add_table_of_contents(doc)
    
    # Introduction
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph(
        'This document specifies the functional and non-functional requirements for the '
        'Codebase Analysis Agent system. The system is designed to analyze legacy codebases '
        'across multiple languages, create dependency graphs, document services, and perform '
        'impact analysis using AI-powered agents.'
    )
    
    # Functional Requirements
    doc.add_heading('2. Functional Requirements', 1)
    
    doc.add_heading('2.1 Repository Management', 2)
    reqs = [
        ('FR-1.1.1', 
         'System SHALL support analyzing repositories from multiple sources: GitHub repositories via API, local file system paths, and Git clone from URLs'),
        ('FR-1.1.2', 'System SHALL support multiple branches per repository'),
        ('FR-1.1.3', 'System SHALL track repository metadata (name, URL, branch, status)')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    doc.add_heading('2.2 Code Parsing', 2)
    reqs = [
        ('FR-1.2.1', 
         'System SHALL parse code in multiple languages: Python (using AST), JavaScript/TypeScript (using regex patterns), Java (using regex patterns), and be extensible to other languages'),
        ('FR-1.2.2', 
         'System SHALL extract: functions, classes, methods, import statements, API endpoint definitions, database connection patterns, and message queue connections')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    doc.add_heading('2.3 Dependency Analysis', 2)
    reqs = [
        ('FR-1.3.1', 'System SHALL identify service boundaries automatically'),
        ('FR-1.3.2', 
         'System SHALL map dependencies between services: direct service-to-service dependencies, API endpoint dependencies, database dependencies, and message queue dependencies'),
        ('FR-1.3.3', 'System SHALL build and maintain a dependency graph in Neo4j')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    doc.add_heading('2.4 Impact Analysis', 2)
    reqs = [
        ('FR-1.4.1', 
         'System SHALL analyze impact of proposed changes: identify affected services, calculate impact scores (0-1), determine risk levels (low, medium, high, critical), and generate recommendations'),
        ('FR-1.4.2', 
         'System SHALL support impact analysis on: specific files, specific services, and change descriptions')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    doc.add_heading('2.5 Documentation Generation', 2)
    reqs = [
        ('FR-1.5.1', 
         'System SHALL generate service documentation using AI: service descriptions, functionality overview, key components, dependencies, and API endpoints'),
        ('FR-1.5.2', 'System SHALL maintain documentation versioning')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    doc.add_heading('2.6 Human-in-the-Loop', 2)
    reqs = [
        ('FR-1.6.1', 
         'System SHALL identify ambiguous cases requiring human input: unclear service boundaries, ambiguous dependencies, and legacy code patterns'),
        ('FR-1.6.2', 'System SHALL pause workflow when checkpoints are created'),
        ('FR-1.6.3', 'System SHALL resume workflow after checkpoint resolution')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    doc.add_heading('2.7 Visualization', 2)
    reqs = [
        ('FR-1.7.1', 'System SHALL provide interactive dependency graph visualization'),
        ('FR-1.7.2', 'System SHALL display service inventory with documentation'),
        ('FR-1.7.3', 'System SHALL show impact analysis results with risk indicators'),
        ('FR-1.7.4', 'System SHALL display agent execution status in real-time')
    ]
    for req_id, req_text in reqs:
        p = doc.add_paragraph()
        p.add_run(f'{req_id}: ').bold = True
        p.add_run(req_text)
    
    # Non-Functional Requirements
    doc.add_heading('3. Non-Functional Requirements', 1)
    
    doc.add_heading('3.1 Performance', 2)
    nfrs = [
        ('NFR-2.1.1', 'System SHALL process repositories up to 10,000 files within 30 minutes'),
        ('NFR-2.1.2', 'API endpoints SHALL respond within 2 seconds for cached data'),
        ('NFR-2.1.3', 'Dependency graph queries SHALL complete within 5 seconds')
    ]
    for nfr_id, nfr_text in nfrs:
        p = doc.add_paragraph()
        p.add_run(f'{nfr_id}: ').bold = True
        p.add_run(nfr_text)
    
    doc.add_heading('3.2 Scalability', 2)
    nfrs = [
        ('NFR-2.2.1', 'System SHALL support horizontal scaling of worker instances'),
        ('NFR-2.2.2', 'System SHALL handle concurrent analysis of multiple repositories'),
        ('NFR-2.2.3', 'Database connections SHALL use connection pooling')
    ]
    for nfr_id, nfr_text in nfrs:
        p = doc.add_paragraph()
        p.add_run(f'{nfr_id}: ').bold = True
        p.add_run(nfr_text)
    
    doc.add_heading('3.3 Reliability', 2)
    nfrs = [
        ('NFR-2.3.1', 'System SHALL have 99% uptime'),
        ('NFR-2.3.2', 'System SHALL gracefully handle agent execution failures'),
        ('NFR-2.3.3', 'System SHALL persist analysis state for recovery')
    ]
    for nfr_id, nfr_text in nfrs:
        p = doc.add_paragraph()
        p.add_run(f'{nfr_id}: ').bold = True
        p.add_run(nfr_text)
    
    doc.add_heading('3.4 Security', 2)
    nfrs = [
        ('NFR-2.4.1', 'System SHALL require API key authentication'),
        ('NFR-2.4.2', 'System SHALL implement rate limiting (60 requests/minute default)'),
        ('NFR-2.4.3', 'System SHALL securely store credentials (GitHub tokens, API keys)'),
        ('NFR-2.4.4', 'System SHALL validate and sanitize all inputs')
    ]
    for nfr_id, nfr_text in nfrs:
        p = doc.add_paragraph()
        p.add_run(f'{nfr_id}: ').bold = True
        p.add_run(nfr_text)
    
    doc.add_heading('3.5 Observability', 2)
    nfrs = [
        ('NFR-2.5.1', 'System SHALL provide structured logging (JSON format)'),
        ('NFR-2.5.2', 'System SHALL expose Prometheus metrics'),
        ('NFR-2.5.3', 'System SHALL track agent execution metrics'),
        ('NFR-2.5.4', 'System SHALL log all errors with context')
    ]
    for nfr_id, nfr_text in nfrs:
        p = doc.add_paragraph()
        p.add_run(f'{nfr_id}: ').bold = True
        p.add_run(nfr_text)
    
    doc.add_heading('3.6 Usability', 2)
    nfrs = [
        ('NFR-2.6.1', 'System SHALL provide RESTful API with OpenAPI documentation'),
        ('NFR-2.6.2', 'System SHALL provide intuitive web interface'),
        ('NFR-2.6.3', 'System SHALL display clear error messages')
    ]
    for nfr_id, nfr_text in nfrs:
        p = doc.add_paragraph()
        p.add_run(f'{nfr_id}: ').bold = True
        p.add_run(nfr_text)
    
    # Technical Requirements
    doc.add_heading('4. Technical Requirements', 1)
    
    doc.add_heading('4.1 Backend', 2)
    doc.add_paragraph('• FastAPI 0.109+', style='List Bullet')
    doc.add_paragraph('• Python 3.10+', style='List Bullet')
    doc.add_paragraph('• Async processing support', style='List Bullet')
    doc.add_paragraph('• Background task processing', style='List Bullet')
    
    doc.add_heading('4.2 Frontend', 2)
    doc.add_paragraph('• Next.js 14+ with App Router', style='List Bullet')
    doc.add_paragraph('• React 18+', style='List Bullet')
    doc.add_paragraph('• TypeScript support', style='List Bullet')
    doc.add_paragraph('• Responsive design', style='List Bullet')
    
    doc.add_heading('4.3 Databases', 2)
    doc.add_paragraph('• Neo4j 5.23+ (graph database)', style='List Bullet')
    doc.add_paragraph('• PostgreSQL 16+ (relational database)', style='List Bullet')
    doc.add_paragraph('• Qdrant (vector database for future semantic search)', style='List Bullet')
    doc.add_paragraph('• Redis 7+ (caching and task queue)', style='List Bullet')
    
    doc.add_heading('4.4 External Services', 2)
    doc.add_paragraph('• OpenAI API (for documentation generation)', style='List Bullet')
    doc.add_paragraph('• GitHub API (for repository access)', style='List Bullet')
    
    # Integration Requirements
    doc.add_heading('5. Integration Requirements', 1)
    
    doc.add_heading('5.1 GitHub Integration', 2)
    doc.add_paragraph('• Support GitHub API v3/v4', style='List Bullet')
    doc.add_paragraph('• Handle authentication tokens', style='List Bullet')
    doc.add_paragraph('• Support private repositories', style='List Bullet')
    
    doc.add_heading('5.2 OpenAI Integration', 2)
    doc.add_paragraph('• Support GPT-4, GPT-4-turbo, GPT-4o-mini', style='List Bullet')
    doc.add_paragraph('• Handle API rate limits', style='List Bullet')
    doc.add_paragraph('• Fallback for API failures', style='List Bullet')
    
    # Data Requirements
    doc.add_heading('6. Data Requirements', 1)
    
    doc.add_heading('6.1 Data Storage', 2)
    doc.add_paragraph('• Repository metadata in PostgreSQL', style='List Bullet')
    doc.add_paragraph('• Dependency graphs in Neo4j', style='List Bullet')
    doc.add_paragraph('• Cached parsed code in Redis', style='List Bullet')
    doc.add_paragraph('• Generated documentation in PostgreSQL', style='List Bullet')
    
    doc.add_heading('6.2 Data Retention', 2)
    doc.add_paragraph('• Analysis results retained for 90 days (configurable)', style='List Bullet')
    doc.add_paragraph('• Cached data expires after 24 hours', style='List Bullet')
    doc.add_paragraph('• Graph data retained indefinitely', style='List Bullet')
    
    # Compliance
    doc.add_heading('7. Compliance Requirements', 1)
    doc.add_paragraph('• Follow OWASP security best practices', style='List Bullet')
    doc.add_paragraph('• GDPR compliance for any user data', style='List Bullet')
    doc.add_paragraph('• Secure credential storage', style='List Bullet')
    
    # Save document
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 
                               'docs', 'Requirements_Document.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

def create_design_document():
    """Create Design Document."""
    doc = Document()
    setup_document_style(doc)
    
    # Title Page
    add_title_page(doc, "Design Document", 
                   "Codebase Analysis Agent System", "1.0.0")
    
    # Table of Contents
    add_table_of_contents(doc)
    
    # Introduction
    doc.add_heading('1. System Architecture', 1)
    
    doc.add_heading('1.1 High-Level Architecture', 2)
    doc.add_paragraph(
        'The system follows a layered architecture with clear separation of concerns:'
    )
    doc.add_paragraph('• Frontend Layer: Next.js 14+ with React, TypeScript, and Tailwind CSS', style='List Bullet')
    doc.add_paragraph('• API Gateway Layer: FastAPI with CORS, rate limiting, and authentication', style='List Bullet')
    doc.add_paragraph('• Business Logic Layer: Agent Orchestrator and Core Services', style='List Bullet')
    doc.add_paragraph('• Data Layer: Neo4j, PostgreSQL, and Redis', style='List Bullet')
    
    doc.add_heading('1.2 Component Architecture', 2)
    
    doc.add_heading('Backend Components', 3)
    components = [
        ('API Layer (backend/api/)', 
         'RESTful endpoints, middleware for rate limiting/CORS/authentication, Pydantic models for validation'),
        ('Agent Layer (backend/agents/)', 
         'Base Agent abstract class, specialized agents for each analysis task, AgentState for execution context'),
        ('Service Layer (backend/services/)', 
         'Repository Manager, Code Parser, Dependency Analyzer, Graph Service, Impact Engine, Cache Service, Agent Orchestrator'),
        ('Parser Layer (backend/parsers/)', 
         'Base Parser interface, language-specific parsers for Python, JavaScript, Java'),
        ('Core Layer (backend/core/)', 
         'Configuration management, database connections, security, logging, monitoring')
    ]
    for comp_name, comp_desc in components:
        p = doc.add_paragraph()
        p.add_run(f'{comp_name}: ').bold = True
        p.add_run(comp_desc)
    
    doc.add_heading('Frontend Components', 3)
    frontend_components = [
        ('Pages (frontend/app/)', 
         'Dashboard, Dependency Graph, Services, Impact Analysis, Agent Status'),
        ('Components (frontend/components/)', 
         'Reusable React components, UI components, Graph components with React Flow'),
        ('Lib (frontend/lib/)', 
         'API client based on Axios, utility functions')
    ]
    for comp_name, comp_desc in frontend_components:
        p = doc.add_paragraph()
        p.add_run(f'{comp_name}: ').bold = True
        p.add_run(comp_desc)
    
    doc.add_heading('1.3 Data Model', 2)
    
    doc.add_heading('Neo4j Graph Model', 3)
    doc.add_paragraph('The graph database stores the following relationships:')
    relationships = [
        '(Service)-[:DEPENDS_ON]->(Service)',
        '(Service)-[:CALLS_API]->(APIEndpoint)',
        '(Service)-[:USES_DB]->(Database)',
        '(Service)-[:PUBLISHES_TO]->(MessageQueue)',
        '(File)-[:BELONGS_TO]->(Service)',
        '(Function)-[:DEFINED_IN]->(File)',
        '(Function)-[:CALLS]->(Function)'
    ]
    for rel in relationships:
        doc.add_paragraph(f'• {rel}', style='List Bullet')
    
    doc.add_heading('PostgreSQL Schema', 3)
    tables = [
        'repositories: Repository metadata',
        'analysis_runs: Analysis execution history',
        'services: Service inventory',
        'documentation: Generated documentation',
        'impact_analyses: Impact analysis results',
        'human_reviews: Human review checkpoints'
    ]
    for table in tables:
        doc.add_paragraph(f'• {table}', style='List Bullet')
    
    # Design Patterns
    doc.add_heading('2. Design Patterns', 1)
    
    patterns = [
        ('Agent Pattern', 
         'Each agent is a specialized component with single responsibility. Agents communicate through shared state (AgentState). Orchestrator coordinates agent execution sequence.'),
        ('Service Layer Pattern', 
         'Business logic separated from API layer. Services are stateless and reusable. Clear separation of concerns.'),
        ('Repository Pattern', 
         'Database access abstracted through services. Models define data structure. Easy to swap database implementations.'),
        ('Strategy Pattern', 
         'Parser selection based on file type. Different parsing strategies for different languages.'),
        ('Observer Pattern (Future)', 
         'Real-time status updates via WebSockets. Event-driven architecture for agent completion.')
    ]
    for pattern_name, pattern_desc in patterns:
        p = doc.add_paragraph()
        p.add_run(f'{pattern_name}: ').bold = True
        p.add_run(pattern_desc)
    
    # API Design
    doc.add_heading('3. API Design', 1)
    
    doc.add_heading('3.1 RESTful Endpoints', 2)
    
    endpoints = [
        ('Repositories', [
            'POST /api/repositories/analyze - Start repository analysis',
            'GET /api/repositories/{id}/status - Get analysis status',
            'GET /api/repositories/ - List all repositories'
        ]),
        ('Services', [
            'GET /api/services/ - List all services',
            'GET /api/services/{id} - Get service details',
            'GET /api/services/{id}/dependencies - Get service dependencies'
        ]),
        ('Dependencies', [
            'GET /api/dependencies/graph - Get dependency graph'
        ]),
        ('Impact Analysis', [
            'POST /api/impact-analysis/analyze - Run impact analysis',
            'GET /api/impact-analysis/{id} - Get impact analysis results'
        ]),
        ('Documentation', [
            'GET /api/documentation/{service_id} - Get service documentation',
            'POST /api/documentation/{service_id}/regenerate - Regenerate docs'
        ]),
        ('Human Review', [
            'GET /api/human-review/checkpoints - List checkpoints',
            'GET /api/human-review/checkpoints/{id} - Get checkpoint details',
            'POST /api/human-review/checkpoints/{id}/resolve - Resolve checkpoint'
        ]),
        ('Metrics', [
            'GET /metrics - Prometheus metrics'
        ])
    ]
    
    for category, endpoint_list in endpoints:
        doc.add_heading(category, 3)
        for endpoint in endpoint_list:
            doc.add_paragraph(f'• {endpoint}', style='List Bullet')
    
    doc.add_heading('3.2 Request/Response Models', 2)
    doc.add_paragraph(
        'All endpoints use Pydantic models for validation. Request models validate input data, '
        'response models ensure consistent output format, and error models provide standardized error responses.'
    )
    
    # Security Design
    doc.add_heading('4. Security Design', 1)
    
    doc.add_heading('4.1 Authentication', 2)
    doc.add_paragraph('• API key authentication via header (X-API-Key)', style='List Bullet')
    doc.add_paragraph('• Configurable API keys per environment', style='List Bullet')
    doc.add_paragraph('• Secure key storage (environment variables)', style='List Bullet')
    
    doc.add_heading('4.2 Authorization', 2)
    doc.add_paragraph('• Role-based access control (future enhancement)', style='List Bullet')
    doc.add_paragraph('• Repository-level permissions (future enhancement)', style='List Bullet')
    
    doc.add_heading('4.3 Input Validation', 2)
    doc.add_paragraph('• Pydantic models validate all inputs', style='List Bullet')
    doc.add_paragraph('• File path sanitization', style='List Bullet')
    doc.add_paragraph('• SQL injection prevention (parameterized queries)', style='List Bullet')
    doc.add_paragraph('• XSS prevention (input sanitization)', style='List Bullet')
    
    doc.add_heading('4.4 Rate Limiting', 2)
    doc.add_paragraph('• Redis-based rate limiting', style='List Bullet')
    doc.add_paragraph('• Configurable limits per endpoint', style='List Bullet')
    doc.add_paragraph('• Per-IP tracking', style='List Bullet')
    
    # Performance Design
    doc.add_heading('5. Performance Design', 1)
    
    doc.add_heading('5.1 Caching Strategy', 2)
    doc.add_paragraph('• Redis caching for parsed code (24-hour TTL)', style='List Bullet')
    doc.add_paragraph('• Dependency graph caching (1-hour TTL)', style='List Bullet')
    doc.add_paragraph('• Cache invalidation on repository updates', style='List Bullet')
    
    doc.add_heading('5.2 Async Processing', 2)
    doc.add_paragraph('• Background tasks for long-running analyses', style='List Bullet')
    doc.add_paragraph('• Non-blocking API responses', style='List Bullet')
    doc.add_paragraph('• Progress tracking via status endpoints', style='List Bullet')
    
    doc.add_heading('5.3 Database Optimization', 2)
    doc.add_paragraph('• Connection pooling for all databases', style='List Bullet')
    doc.add_paragraph('• Indexed queries in Neo4j', style='List Bullet')
    doc.add_paragraph('• Efficient graph traversal algorithms', style='List Bullet')
    
    doc.add_heading('5.4 Scalability', 2)
    doc.add_paragraph('• Horizontal scaling support', style='List Bullet')
    doc.add_paragraph('• Stateless service design', style='List Bullet')
    doc.add_paragraph('• Queue-based task processing (future)', style='List Bullet')
    
    # Error Handling
    doc.add_heading('6. Error Handling Design', 1)
    
    doc.add_heading('6.1 Error Types', 2)
    error_types = [
        'Validation errors (400)',
        'Authentication errors (401)',
        'Not found errors (404)',
        'Server errors (500)'
    ]
    for error_type in error_types:
        doc.add_paragraph(f'• {error_type}', style='List Bullet')
    
    doc.add_heading('6.2 Error Response Format', 2)
    doc.add_paragraph('Errors follow a standardized JSON format:')
    doc.add_paragraph('{\n  "detail": "Error message",\n  "error_code": "ERROR_CODE",\n  "timestamp": "2024-01-01T00:00:00Z"\n}', 
                     style='No Spacing')
    
    doc.add_heading('6.3 Logging Strategy', 2)
    doc.add_paragraph('• Structured JSON logging', style='List Bullet')
    doc.add_paragraph('• Error context preservation', style='List Bullet')
    doc.add_paragraph('• Log levels: DEBUG, INFO, WARNING, ERROR, CRITICAL', style='List Bullet')
    
    # Monitoring
    doc.add_heading('7. Monitoring Design', 1)
    
    doc.add_heading('7.1 Metrics', 2)
    metrics = [
        'Request count and duration',
        'Agent execution metrics',
        'Active analysis count',
        'Cache hit/miss rates',
        'Database query performance'
    ]
    for metric in metrics:
        doc.add_paragraph(f'• {metric}', style='List Bullet')
    
    doc.add_heading('7.2 Logging', 2)
    doc.add_paragraph('• Structured logging (JSON)', style='List Bullet')
    doc.add_paragraph('• Request/response logging', style='List Bullet')
    doc.add_paragraph('• Agent execution logging', style='List Bullet')
    doc.add_paragraph('• Error stack traces', style='List Bullet')
    
    doc.add_heading('7.3 Health Checks', 2)
    doc.add_paragraph('• /health endpoint', style='List Bullet')
    doc.add_paragraph('• Database connectivity checks', style='List Bullet')
    doc.add_paragraph('• External service availability', style='List Bullet')
    
    # Deployment
    doc.add_heading('8. Deployment Design', 1)
    
    doc.add_heading('8.1 Containerization', 2)
    doc.add_paragraph('• Docker containers for all services', style='List Bullet')
    doc.add_paragraph('• Docker Compose for local development', style='List Bullet')
    doc.add_paragraph('• Multi-stage builds for production', style='List Bullet')
    
    doc.add_heading('8.2 Environment Configuration', 2)
    doc.add_paragraph('• Environment variables for all configuration', style='List Bullet')
    doc.add_paragraph('• Separate configs for dev/staging/prod', style='List Bullet')
    doc.add_paragraph('• Secret management (future: Vault integration)', style='List Bullet')
    
    doc.add_heading('8.3 Database Migrations', 2)
    doc.add_paragraph('• Alembic for PostgreSQL migrations', style='List Bullet')
    doc.add_paragraph('• Version-controlled schema changes', style='List Bullet')
    doc.add_paragraph('• Rollback support', style='List Bullet')
    
    # Future Enhancements
    doc.add_heading('9. Future Enhancements', 1)
    
    doc.add_heading('9.1 Planned Features', 2)
    features = [
        'WebSocket support for real-time updates',
        'Advanced semantic search using Qdrant',
        'Multi-repository comparison',
        'Change history tracking',
        'Automated testing recommendations'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('9.2 Scalability Improvements', 2)
    improvements = [
        'Message queue (RabbitMQ/Kafka) for task processing',
        'Distributed agent execution',
        'Graph database sharding',
        'CDN for frontend assets'
    ]
    for improvement in improvements:
        doc.add_paragraph(f'• {improvement}', style='List Bullet')
    
    doc.add_heading('9.3 AI Enhancements', 2)
    ai_features = [
        'Code quality scoring',
        'Automated refactoring suggestions',
        'Security vulnerability detection',
        'Performance bottleneck identification'
    ]
    for feature in ai_features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    # Save document
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 
                               'docs', 'Design_Document.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

if __name__ == '__main__':
    print("Generating documentation...")
    print("=" * 50)
    create_process_flow_document()
    create_requirements_document()
    create_design_document()
    print("=" * 50)
    print("All documents created successfully!")
